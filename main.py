#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2021/5/30 13:16
# @Author  : oreki
# @File    : main.py
# @Software: PyCharm
# @Mail    : a912550157@gmail.com
import json
import os
import pymysql
import re
import docx
from get import GetByCDN
import logging

logging.basicConfig(level=logging.DEBUG,  # 控制台打印的日志级别
                    filename='new.log',
                    filemode='a',  ##模式，有w和a，w就是写模式，每次都会重新写日志，覆盖之前的日志
                    # a是追加模式，默认如果不写的话，就是追加模式
                    format=
                    '%(asctime)s - %(pathname)s[line:%(lineno)d] - %(levelname)s: %(message)s'
                    # 日志格式
                    )

CDN = GetByCDN()

info = {
    "host": "",
    "user": "",
    "password": "",
    "path": ""
}


def db_operation():
    db = pymysql.connect(host=info['host'], port=3306, user=info['user'], passwd=info['password'], db='huochai',
                         charset="utf8")
    cursor = db.cursor()
    sql = "select id, units from Msg where AuthorId = 504 and status = 1"
    cursor.execute(sql)
    results = cursor.fetchall()
    file = docx.Document()  # 创建内存中的word文档对象
    a = 0
    size = 0
    for row in results:
        unit = row[1]
        unit = json.loads(unit)
        content = CDN.getdata(units=unit)
        logging.info("start..{}".format(row[0]))
        for i in content:
            file.add_paragraph("----------------------------------------")
            if "https:" in i:
                path = CDN.download_pic(i)
                img_szie = get_size(path)
                file.add_picture(path)
                os.remove(path)
                size += img_szie
            else:
                file.add_paragraph(u"{}".format(i))  # 写入若干段落
            # print("size data", size)
            if size > 10:
                file.save("{}/{}.docx".format(info['path'], a))  # 保存才能看到结果
                size = 0
                a += 1
                file = docx.Document()
            else:
                file.save("{}/{}.docx".format(info['path'], a))  # 保存才能看到结果


def get_size(filepath):
    fsize = os.path.getsize(filepath)
    fsize = fsize / float(1024 * 1024)
    return round(fsize, 2)


if __name__ == '__main__':
    try:
        db_operation()
    finally:
        logging.shutdown()
